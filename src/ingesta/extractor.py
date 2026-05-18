"""Extractor de texto de documentos (PDF, Word, imagen)."""
import os
from dataclasses import dataclass
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document


@dataclass
class DocumentoExtraido:
    """Resultado de extracción de texto de un documento."""
    texto: str           # texto extraído (vacío si es imagen sin OCR)
    paginas: int         # número de páginas (1 para Word/imagen)
    es_imagen: bool      # True si el doc es plano/imagen (no texto extraíble)
    bytes_raw: bytes     # bytes originales del archivo (para guardar en Storage)
    paginas_texto: list[str] = None  # PDF: una entrada por página (índice 0 = pág 1)


def extraer(ruta_archivo: str) -> DocumentoExtraido:
    """
    Extrae texto de un documento según su extensión.
    - .pdf  → PyMuPDF (fitz). Si el PDF tiene texto extraíble, devuelve texto.
              Si todas las páginas están vacías (PDF de imagen), setear es_imagen=True.
    - .docx → python-docx. Extrae párrafos y texto de tablas.
    - .doc  → No suportado: lanzar ValueError con mensaje claro.
    - .png / .jpg / .jpeg → es_imagen=True, texto="" (OCR fuera del MVP).
    - otros → lanzar ValueError(f"Tipo de archivo no soportado: {ext}")
    """
    ruta = Path(ruta_archivo)
    ext = ruta.suffix.lower()

    if ext not in ['.pdf', '.docx', '.doc', '.png', '.jpg', '.jpeg']:
        raise ValueError(f"Tipo de archivo no soportado: {ext}")

    if ext == '.doc':
        raise ValueError("Formato .doc no soportado. Usá .docx")

    # Leer bytes crudos
    with open(ruta, 'rb') as f:
        bytes_raw = f.read()

    if ext == '.pdf':
        return _extraer_pdf(bytes_raw)
    elif ext == '.docx':
        return _extraer_docx(bytes_raw)
    else:
        # Imagen
        return DocumentoExtraido(
            texto="",
            paginas=1,
            es_imagen=True,
            bytes_raw=bytes_raw,
            paginas_texto=[],
        )


def _extraer_pdf(bytes_raw: bytes) -> DocumentoExtraido:
    """Extrae texto de PDF con PyMuPDF."""
    doc = fitz.open(stream=bytes_raw, filetype="pdf")
    num_paginas = len(doc)

    if num_paginas == 0:
        return DocumentoExtraido(
            texto="",
            paginas=0,
            es_imagen=True,
            bytes_raw=bytes_raw,
            paginas_texto=[],
        )

    textos_paginas = []  # Una entrada por página (índice 0 = página 1)
    textos = []        # Texto concatenado con marcadores
    tiene_texto = False

    for i, pagina in enumerate(doc, start=1):
        texto = pagina.get_text().strip()
        if texto:
            tiene_texto = True
            textos_paginas.append(texto)
            textos.append(f"\n\n--- Página {i} ---\n\n{texto}")
        else:
            textos_paginas.append("")
            textos.append(f"\n\n--- Página {i} ---\n\n")

    doc.close()

    texto_final = "".join(textos).strip()
    # paginas_texto: índice 0 = página 1
    paginas_texto = textos_paginas

    return DocumentoExtraido(
        texto=texto_final,
        paginas=num_paginas,
        es_imagen=not tiene_texto,
        bytes_raw=bytes_raw,
        paginas_texto=paginas_texto,
    )


def _extraer_docx(bytes_raw: bytes) -> DocumentoExtraido:
    """Extrae texto de Word con python-docx."""
    import io
    doc = Document(io.BytesIO(bytes_raw))

    parrafos = [p.text for p in doc.paragraphs if p.text.strip()]

    # Texto de tablas
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                if celda.text.strip():
                    parrafos.append(celda.text.strip())

    texto = "\n".join(parrafos)

    return DocumentoExtraido(
        texto=texto,
        paginas=1,
        es_imagen=False,
        bytes_raw=bytes_raw,
        paginas_texto=[],
    )